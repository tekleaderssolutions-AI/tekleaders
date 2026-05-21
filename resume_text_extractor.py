"""Extract plain text from resume files: PDF, Word (.docx/.doc), and ZIP archives."""
from __future__ import annotations

import sys
import tempfile
import zipfile
from io import BytesIO
from pathlib import PurePosixPath
from typing import List, Tuple

import pdfplumber

RESUME_EXTENSIONS = frozenset({".pdf", ".doc", ".docx"})
ZIP_EXTENSIONS = frozenset({".zip"})
SUPPORTED_UPLOAD_EXTENSIONS = RESUME_EXTENSIONS | ZIP_EXTENSIONS

MAX_ZIP_ENTRIES = 50
MAX_ENTRY_BYTES = 15 * 1024 * 1024


def _normalize_name(name: str) -> str:
    return (name or "").replace("\\", "/").strip()


def _ext(name: str) -> str:
    return PurePosixPath(_normalize_name(name)).suffix.lower()


def is_supported_upload_filename(filename: str) -> bool:
    return _ext(filename or "") in SUPPORTED_UPLOAD_EXTENSIONS


def is_resume_filename(filename: str) -> bool:
    return _ext(filename or "") in RESUME_EXTENSIONS


def expand_upload(filename: str, contents: bytes) -> List[Tuple[str, bytes]]:
    """
    Expand one upload into (display_name, bytes) pairs.
    ZIP archives yield one entry per supported resume inside.
    """
    if not contents:
        return []

    ext = _ext(filename)
    if ext in ZIP_EXTENSIONS:
        return _expand_zip(filename, contents)
    if ext in RESUME_EXTENSIONS:
        return [(filename, contents)]
    return []


def _expand_zip(archive_name: str, contents: bytes) -> List[Tuple[str, bytes]]:
    out: List[Tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(BytesIO(contents)) as zf:
            all_entries = [
                n
                for n in zf.namelist()
                if n and not n.endswith("/") and not n.endswith("\\")
                and not _normalize_name(n).startswith("__MACOSX/")
                and not PurePosixPath(_normalize_name(n)).name.startswith(".")
            ]
            names = [n for n in all_entries if is_resume_filename(n)]
            if not names:
                if not all_entries:
                    raise ValueError("ZIP archive is empty")
                inside = ", ".join(
                    PurePosixPath(_normalize_name(n)).name for n in all_entries[:12]
                )
                raise ValueError(
                    "No PDF or Word files inside ZIP. "
                    f"Add .pdf, .doc, or .docx files (found: {inside})"
                )
            if len(names) > MAX_ZIP_ENTRIES:
                raise ValueError(f"ZIP contains more than {MAX_ZIP_ENTRIES} resume files")
            for name in names:
                info = zf.getinfo(name)
                if info.file_size > MAX_ENTRY_BYTES:
                    raise ValueError(f"File too large in ZIP: {name}")
                data = zf.read(name)
                display = f"{archive_name}::{PurePosixPath(_normalize_name(name)).name}"
                out.append((display, data))
    except zipfile.BadZipFile as e:
        raise ValueError("Invalid or corrupted ZIP file") from e
    return out


def extract_resume_text(filename: str, contents: bytes) -> str:
    ext = _ext(filename)
    if ext == ".pdf":
        return _extract_pdf_text(contents)
    if ext == ".docx":
        return _extract_docx_text(contents)
    if ext == ".doc":
        return _extract_doc_text(contents)
    raise ValueError(f"Unsupported resume format: {ext or '(no extension)'}")


def _extract_pdf_text(contents: bytes) -> str:
    with pdfplumber.open(BytesIO(contents)) as pdf:
        return "\n".join([page.extract_text() or "" for page in pdf.pages])


def _extract_docx_text(contents: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("python-docx is not installed. Run: pip install python-docx") from e

    doc = Document(BytesIO(contents))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_doc_text(contents: bytes) -> str:
    """Legacy .doc — try docx parser if file is OOXML, else Word COM on Windows."""
    if contents[:2] == b"PK":
        return _extract_docx_text(contents)
    if sys.platform == "win32":
        text = _extract_doc_via_word_com(contents)
        if text and text.strip():
            return text
    raise ValueError(
        "Could not read legacy .doc file. Open in Word and Save As .docx or PDF, then re-upload."
    )


def _extract_doc_via_word_com(contents: bytes) -> str:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return ""

    tmp_path = None
    word = None
    doc = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(contents)
            tmp_path = tmp.name

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(tmp_path, ReadOnly=True)
        return str(doc.Content.Text or "")
    except Exception:
        return ""
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if tmp_path:
            try:
                import os

                os.unlink(tmp_path)
            except OSError:
                pass
