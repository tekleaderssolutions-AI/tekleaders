"""Extract plain text from job description uploads (PDF, Word)."""
from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath

import pdfplumber

JD_EXTENSIONS = frozenset({".pdf", ".docx", ".doc"})


def jd_extension(filename: str) -> str:
    return PurePosixPath((filename or "").replace("\\", "/")).suffix.lower()


def is_supported_jd_filename(filename: str) -> bool:
    return jd_extension(filename) in JD_EXTENSIONS


def extract_jd_text(filename: str, contents: bytes) -> str:
    ext = jd_extension(filename)
    if ext == ".pdf":
        return _extract_pdf(contents)
    if ext == ".docx":
        return _extract_docx(contents)
    if ext == ".doc":
        return _extract_doc(contents)
    raise ValueError(f"Unsupported JD format: {ext or '(no extension)'}. Use PDF or Word (.docx).")


def _extract_pdf(contents: bytes) -> str:
    with pdfplumber.open(BytesIO(contents)) as pdf:
        return "\n".join([page.extract_text() or "" for page in pdf.pages])


def _extract_docx(contents: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("python-docx is not installed. Run: pip install python-docx") from e

    doc = Document(BytesIO(contents))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_doc(contents: bytes) -> str:
    if contents[:2] == b"PK":
        return _extract_docx(contents)
    raise ValueError(
        "Legacy .doc is not supported for JD upload. Open in Word and Save As .docx or PDF."
    )
